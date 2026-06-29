from __future__ import annotations

import html
import math
import re
import shutil
import textwrap
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\ACER\Downloads\Course Project Report Template.docx")
DRAWIO = Path(r"C:\Users\ACER\Downloads\Topic4_SWD_Group3.drawio.xml")
OUT_DIR = ROOT / "reports"
ASSET_DIR = OUT_DIR / "seal_diagrams"
OUT_DOCX = OUT_DIR / "SEAL_Course_Project_Report.docx"


def strip_label(value: str | None) -> str:
    if not value:
        return ""
    text = html.unescape(value)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"</(div|p|li|tr)>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = html.unescape(text)
    text = text.replace("\xa0", " ")
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def parse_style(style: str | None) -> dict[str, str]:
    result: dict[str, str] = {}
    if not style:
        return result
    for part in style.split(";"):
        if "=" in part:
            key, val = part.split("=", 1)
            result[key] = val
        elif part:
            result[part] = "1"
    return result


def color(style: dict[str, str], key: str, default: str) -> str:
    raw = style.get(key, default)
    if raw.startswith("#") and len(raw) == 7:
        return raw
    if raw in ("none", "transparent"):
        return default
    if raw.startswith("light-dark("):
        m = re.search(r"#([0-9a-fA-F]{6})", raw)
        if m:
            return "#" + m.group(1)
    names = {
        "white": "#ffffff",
        "black": "#000000",
        "gray": "#808080",
        "lightgray": "#d9d9d9",
    }
    return names.get(raw.lower(), default)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def fit_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, width: int, max_lines: int = 8) -> list[str]:
    if not text:
        return []
    lines: list[str] = []
    for original in text.splitlines():
        words = original.split()
        current = ""
        for word in words:
            trial = (current + " " + word).strip()
            if draw.textbbox((0, 0), trial, font=font)[2] <= width or not current:
                current = trial
            else:
                lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines[:max_lines]


def diagram_elements(diagram: ET.Element):
    cells: dict[str, dict] = {}
    edges: list[dict] = []
    for el in diagram.iter():
        if el.tag not in {"mxCell", "UserObject"}:
            continue
        cell = el if el.tag == "mxCell" else el.find("mxCell")
        if cell is None:
            continue
        geom = cell.find("mxGeometry")
        if geom is None:
            continue
        cid = el.get("id") or cell.get("id")
        label = strip_label(el.get("label") or cell.get("value") or el.get("mermaidBaseValue"))
        style = parse_style(cell.get("style") or el.get("mermaidBaseStyle"))
        x = float(geom.get("x", "0"))
        y = float(geom.get("y", "0"))
        w = float(geom.get("width", "0"))
        h = float(geom.get("height", "0"))
        if cell.get("vertex") == "1" and w and h:
            cells[cid] = {"id": cid, "label": label, "style": style, "x": x, "y": y, "w": w, "h": h}
        elif cell.get("edge") == "1":
            pts = []
            arr = geom.find("Array")
            if arr is not None:
                for p in arr.findall("mxPoint"):
                    pts.append((float(p.get("x", "0")), float(p.get("y", "0"))))
            edges.append(
                {
                    "label": label,
                    "style": style,
                    "source": cell.get("source"),
                    "target": cell.get("target"),
                    "points": pts,
                    "x": x,
                    "y": y,
                }
            )
    return cells, edges


def render_diagram(diagram: ET.Element, output: Path) -> bool:
    cells, edges = diagram_elements(diagram)
    if not cells:
        return False
    xs = [c["x"] for c in cells.values()] + [c["x"] + c["w"] for c in cells.values()]
    ys = [c["y"] for c in cells.values()] + [c["y"] + c["h"] for c in cells.values()]
    min_x, max_x = min(xs), max(xs)
    min_y, max_y = min(ys), max(ys)
    margin = 60
    raw_w = max_x - min_x + margin * 2
    raw_h = max_y - min_y + margin * 2
    scale = min(2400 / raw_w, 1800 / raw_h, 1.5)
    width, height = int(raw_w * scale), int(raw_h * scale)
    img = Image.new("RGB", (max(width, 800), max(height, 500)), "#ffffff")
    draw = ImageDraw.Draw(img)

    def tx(x): return int((x - min_x + margin) * scale)
    def ty(y): return int((y - min_y + margin) * scale)

    def center(cid: str | None):
        if not cid or cid not in cells:
            return None
        c = cells[cid]
        return tx(c["x"] + c["w"] / 2), ty(c["y"] + c["h"] / 2)

    line_font = load_font(max(10, int(12 * scale)))
    for edge in edges:
        points = [center(edge["source"])]
        points.extend((tx(x), ty(y)) for x, y in edge["points"])
        points.append(center(edge["target"]))
        points = [p for p in points if p]
        if len(points) < 2:
            continue
        stroke = color(edge["style"], "strokeColor", "#555555")
        draw.line(points, fill=stroke, width=max(1, int(2 * scale)))
        x1, y1 = points[-2]
        x2, y2 = points[-1]
        angle = math.atan2(y2 - y1, x2 - x1)
        arrow = 10 * scale
        a1 = angle + math.pi * 0.82
        a2 = angle - math.pi * 0.82
        draw.polygon(
            [(x2, y2), (x2 + arrow * math.cos(a1), y2 + arrow * math.sin(a1)), (x2 + arrow * math.cos(a2), y2 + arrow * math.sin(a2))],
            fill=stroke,
        )
        if edge["label"]:
            mx = sum(p[0] for p in points) // len(points)
            my = sum(p[1] for p in points) // len(points)
            draw.text((mx + 4, my + 4), edge["label"][:60], font=line_font, fill="#333333")

    for c in sorted(cells.values(), key=lambda item: item["w"] * item["h"], reverse=True):
        x, y, w, h = tx(c["x"]), ty(c["y"]), int(c["w"] * scale), int(c["h"] * scale)
        st = c["style"]
        fill = color(st, "fillColor", "#f8fbff")
        stroke = color(st, "strokeColor", "#4f6f91")
        shape = st.get("shape", "")
        rounded = "rounded=1" in ";".join(f"{k}={v}" for k, v in st.items())
        if "ellipse" in shape:
            draw.ellipse([x, y, x + w, y + h], fill=fill, outline=stroke, width=max(1, int(2 * scale)))
        elif "rhombus" in shape:
            draw.polygon([(x + w // 2, y), (x + w, y + h // 2), (x + w // 2, y + h), (x, y + h // 2)], fill=fill, outline=stroke)
        else:
            if rounded:
                draw.rounded_rectangle([x, y, x + w, y + h], radius=max(6, int(10 * scale)), fill=fill, outline=stroke, width=max(1, int(2 * scale)))
            else:
                draw.rectangle([x, y, x + w, y + h], fill=fill, outline=stroke, width=max(1, int(2 * scale)))
        label = c["label"]
        if label:
            font_size = max(10, min(22, int(float(st.get("fontSize", "12")) * scale)))
            font = load_font(font_size, bold=st.get("fontStyle") == "1")
            lines = fit_text(draw, label, font, max(30, w - 12), max_lines=max(2, h // max(12, font_size)))
            total_h = len(lines) * (font_size + 2)
            start_y = y + max(4, (h - total_h) // 2)
            for line in lines:
                bbox = draw.textbbox((0, 0), line, font=font)
                draw.text((x + max(4, (w - (bbox[2] - bbox[0])) // 2), start_y), line, font=font, fill=color(st, "fontColor", "#1f2937"))
                start_y += font_size + 2

    output.parent.mkdir(parents=True, exist_ok=True)
    img.save(output, quality=95)
    return True


def render_all_diagrams() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    tree = ET.parse(DRAWIO)
    images: dict[str, Path] = {}
    for index, diagram in enumerate(tree.getroot().findall("diagram"), start=1):
        name = diagram.get("name", f"Diagram {index}")
        safe = re.sub(r"[^A-Za-z0-9_-]+", "_", name).strip("_")
        path = ASSET_DIR / f"{index:02d}_{safe}.png"
        if render_diagram(diagram, path):
            images[name] = path
    return images


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.1
    for name, size, color_hex, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color_hex)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Course Project Report")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor.from_string("1F4D78")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Subject: SWD392")
    r.font.size = Pt(15)
    r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SEAL - Software Engineering Hackathon Management System")
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("2E74B5")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Vietnamese title: Hệ thống quản lý cuộc thi SEAL Hackathon ngành Kỹ thuật Phần mềm")
    for text in ["Group: SWD392-G3", "Da Nang, June 2026"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).font.size = Pt(12)
    doc.add_page_break()


def add_para(doc: Document, text: str, style: str | None = None, bold_label: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(f"• {item}", style="List Paragraph")


def add_numbered(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        doc.add_paragraph(f"{index}. {item}", style="List Paragraph")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph("")
    return table


def add_figure(doc: Document, images: dict[str, Path], name: str, caption: str, width: float = 6.2) -> None:
    path = images.get(name)
    if not path:
        add_para(doc, f"[Diagram placeholder: {name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)


def build_report(images: dict[str, Path]) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    style_doc(doc)
    add_title(doc)

    doc.add_heading("Record of Changes", level=1)
    add_table(doc, ["Date", "A/M/D", "In charge", "Change Description"], [["29/06/2026", "A", "Group 3", "Initial course project report for SEAL Hackathon Management System."]], [1.2, 0.8, 1.4, 3.1])

    doc.add_heading("I. Overview", level=1)
    doc.add_heading("I.1 Project Information", level=2)
    add_bullets(doc, [
        "Project name: SEAL - Software Engineering Hackathon Management System",
        "Vietnamese title: Hệ thống quản lý cuộc thi SEAL Hackathon ngành Kỹ thuật Phần mềm",
        "Project code: SEAL",
        "Group name: SWD392-G3",
        "Software type: Web application with REST API backend and relational database",
    ])
    doc.add_heading("I.2 Project Team", level=2)
    add_table(doc, ["Full Name", "Role", "Email", "Mobile"], [
        ["Nguyen Trung Kien", "Lecturer", "kiennt@fe.edu.vn", "0912656836"],
        ["TBD", "Leader", "TBD", "TBD"],
        ["TBD", "Member", "TBD", "TBD"],
        ["TBD", "Member", "TBD", "TBD"],
        ["TBD", "Member", "TBD", "TBD"],
    ], [2.0, 1.3, 2.1, 1.1])

    doc.add_heading("II. Requirement Specification", level=1)
    doc.add_heading("II.1 Problem description", level=2)
    add_para(doc, "SEAL is designed to support Software Engineering hackathon operations from event setup to final award announcement. The current process of registering teams, assigning judges, collecting submissions, scoring, ranking, and eliminating violating teams can become fragmented across spreadsheets, forms, email, and manual calculations. This creates risks of inconsistent criteria, delayed results, unclear audit trails, and limited research data for rubric-based learning evaluation.")
    add_para(doc, "The system provides a controlled platform where organizers configure events and rounds, students form teams and submit project links, mentors support categories, judges score submissions by criteria, and the system automatically calculates rankings and promotion results.")

    doc.add_heading("II.2 Major Features", level=2)
    features = [
        ["FE-01", "Event and round management", "Create hackathon events, configure qualification/final rounds, deadlines, judges, scoring criteria, and top-N promotion rules."],
        ["FE-02", "Scoring criteria management", "Maintain reusable default rubrics and allow event-specific criteria/weight adjustments."],
        ["FE-03", "Category and mentor management", "Create event categories and assign mentors; the same lecturer may mentor one category and judge another category."],
        ["FE-04", "Team management", "Create teams with 3-5 members and register a team into one event category."],
        ["FE-05", "User registration and approval", "Support email/password authentication with JWT, student classification, organizer approval, and temporary guest judge accounts."],
        ["FE-06", "Submission management", "Allow teams to submit repository, demo, report, and slide URLs for each round; optionally integrate GitHub/GitLab metadata."],
        ["FE-07", "Evaluation", "Allow assigned judges to score each submission by event criteria while storing every judge/criterion score separately."],
        ["FE-08", "Ranking, promotion, and elimination", "Automatically rank teams by round/category/event, promote qualified teams, and record eliminations with reasons and audit logs."],
        ["FE-09", "RBL data collection", "Record detailed scores, run calibration rounds, export anonymized CSV, and show variance dashboards for inter-rater reliability analysis."],
        ["FE-10", "Awards and reporting", "Publish results, notify participants, and export rankings and score reports to CSV/Excel."],
    ]
    add_table(doc, ["#", "Feature", "System Function", "Description"], [[str(i + 1), f[1], f[0], f[2]] for i, f in enumerate(features)], [0.45, 1.75, 1.15, 3.15])

    doc.add_heading("II.3 Context Diagram", level=2)
    add_para(doc, "The system boundary includes participants, organizers, mentors, internal judges, guest judges, and optional external source-control services. Participants submit project artifacts, organizers configure the competition and approve accounts, judges evaluate submissions, and the system produces rankings, promotion results, audit logs, and exportable reports.")
    add_figure(doc, images, "HighSystem", "Figure 1. Context and high-level system view.")

    doc.add_heading("II.4 Nonfunctional requirements", level=2)
    add_table(doc, ["ID", "Requirement", "Description"], [
        ["NFR-01", "Security", "Use JWT authentication, role-based authorization, password hashing, and organizer approval before competition participation."],
        ["NFR-02", "Auditability", "Record all scoring, elimination, promotion, and administrative changes with actor, time, action, and reason where applicable."],
        ["NFR-03", "Reliability", "Preserve submitted URLs, score details, and ranking snapshots even when external GitHub/GitLab metadata fetching fails."],
        ["NFR-04", "Performance", "Ranking calculation should complete within seconds for typical course-scale events and remain batch-friendly for larger events."],
        ["NFR-05", "Maintainability", "Separate presentation, business, and data access responsibilities to allow feature extension and testing."],
        ["NFR-06", "Exportability", "Reports and research datasets must be exportable to CSV/Excel with anonymization support for RBL use."],
    ], [0.75, 1.55, 4.2])

    doc.add_heading("II.5 Functional requirements", level=2)
    doc.add_heading("II.5.1 Actors", level=3)
    add_table(doc, ["#", "Actor", "Description"], [
        ["1", "Organizer/Admin", "Creates events, rounds, categories, criteria, user approvals, judge assignments, eliminations, and award publication."],
        ["2", "Student/Participant", "Registers an account, forms or joins a team, registers to a category, and submits project artifacts."],
        ["3", "Mentor", "Supports assigned event categories and monitors teams in that category."],
        ["4", "Internal Judge", "Scores assigned submissions according to event criteria."],
        ["5", "Guest Judge", "Uses a temporary organizer-created account to score only assigned rounds."],
        ["6", "External Git Provider", "Optional GitHub/GitLab API source for repository metadata."],
    ], [0.45, 1.65, 4.4])

    doc.add_heading("II.5.2 Use Cases", level=3)
    add_table(doc, ["ID", "Use Case", "Actors", "Use Case Description"], [
        ["UC-01", "Register Team to Track", "Student, Organizer", "Students create teams of 3-5 members and register to one category after account approval."],
        ["UC-02", "Submit Project", "Team Member", "A team submits repository, demo, report, and slide URLs for the current round before the deadline."],
        ["UC-03", "Judge Scores Submission", "Judge, Guest Judge", "Assigned judges score a submission by criteria and the system stores criterion-level score details."],
        ["UC-04", "Generate Ranking and Promote Teams", "Organizer, System", "The system ranks teams and identifies top-N teams per category for the next round."],
        ["UC-05", "Manage Event and Rounds", "Organizer", "Organizer creates the event, configures rounds, deadlines, judges, criteria, categories, and promotion rules."],
    ], [0.75, 1.65, 1.45, 2.65])
    doc.add_heading("II.5.2.1 Diagram(s)", level=3)
    add_figure(doc, images, "UC Diagram", "Figure 2. Use case diagram for SEAL.")

    doc.add_heading("II.5.2.2 Use case descriptions", level=3)
    usecases = [
        ("UC-01 Register Team to Track", "Student", "The student chooses members and a category. The system validates team size, account approval status, and category availability before creating the registration."),
        ("UC-02 Submit Project", "Team Member", "The team member submits required URLs for a round. The system validates deadline and saves the submission; optional repository metadata is fetched asynchronously."),
        ("UC-03 Judge Scores Submission", "Judge/Guest Judge", "The judge opens an assigned submission, scores each criterion, and submits the evaluation. Scores are stored separately by judge and criterion."),
        ("UC-04 Generate Ranking and Promote Teams", "Organizer/System", "The organizer triggers ranking calculation. The system aggregates weighted scores, sorts teams by category and round, and marks top-N teams as promoted."),
        ("UC-05 Manage Event and Rounds", "Organizer", "The organizer configures an event, categories, rounds, criteria, judge assignments, deadlines, and promotion rules."),
    ]
    for uc, actor, desc in usecases:
        add_table(doc, ["Field", "Content"], [
            ["UC ID and Name", uc],
            ["Primary Actor", actor],
            ["Trigger", "The actor starts the use case from the appropriate SEAL management screen."],
            ["Description", desc],
            ["Preconditions", "The actor is authenticated and authorized; the target event/round/category exists."],
            ["Postconditions", "The requested record is saved and relevant audit data is available."],
            ["Exceptions", "Invalid permission, deadline expired, missing required data, or violated business rule."],
            ["Priority", "High"],
        ], [1.6, 4.9])

    doc.add_heading("b. Business Rules", level=4)
    add_table(doc, ["ID", "Business Rule", "Business Rule Description"], [
        ["BR-01", "Team size", "Each competing team must contain from 3 to 5 members."],
        ["BR-02", "Single category registration", "A team registers into one category for a specific event."],
        ["BR-03", "Approval required", "Accounts must be approved by the organizer before participating."],
        ["BR-04", "Guest judge scope", "Guest judge accounts can score only assigned rounds/submissions."],
        ["BR-05", "Promotion rule", "Only top-N teams in each category advance to the next round according to event configuration."],
        ["BR-06", "Elimination audit", "Eliminated teams/submissions must have a recorded reason and canceled result."],
    ], [0.75, 1.65, 4.1])

    doc.add_heading("II.5.3 Activity diagram", level=3)
    add_para(doc, "The activity flow follows the competition lifecycle: organizer configures the event, students register and form teams, teams submit round artifacts, judges evaluate submissions, and the system calculates rankings, promotions, eliminations, and awards.")

    doc.add_heading("III. Analysis models.", level=1)
    doc.add_heading("III.1 Interaction diagrams", level=2)
    for name, caption in [
        ("SD01 - Register Team to Track", "Figure 3. Sequence diagram - Register Team to Track."),
        ("SD02 - Submit Project", "Figure 4. Sequence diagram - Submit Project."),
        ("SD03 - Judge Scores Submission", "Figure 5. Sequence diagram - Judge Scores Submission."),
        ("SD04 - Generate Ranking and Promote Teams", "Figure 6. Sequence diagram - Generate Ranking and Promote Teams."),
        ("SD05 - Manage Event and Rounds", "Figure 7. Sequence diagram - Manage Event and Rounds."),
    ]:
        add_figure(doc, images, name, caption)
    doc.add_heading("III.2 State diagram", level=2)
    add_para(doc, "The main stateful objects are Event, Round, TeamRegistration, Submission, Evaluation, and Award. A team registration starts as Pending/Registered, moves through Submitted and Evaluated per round, and may become Promoted, Eliminated, or Awarded depending on ranking and organizer actions.")

    doc.add_heading("IV. Design specification", level=1)
    doc.add_heading("IV.1 Integrated Communication Diagrams", level=2)
    for name, caption in [
        ("C01_Communication_Diagram", "Figure 8. Communication diagram - team registration and event setup interactions."),
        ("C02_Communication_Diagram", "Figure 9. Communication diagram - submission and judging interactions."),
        ("C03_Communication_Diagram", "Figure 10. Communication diagram - ranking, promotion, and reporting interactions."),
    ]:
        add_figure(doc, images, name, caption)

    doc.add_heading("IV.2 System High-Level Design", level=2)
    add_para(doc, "SEAL uses a layered architecture. The presentation layer exposes web UI and REST API endpoints. The application/service layer coordinates use cases such as event setup, team registration, submission, scoring, ranking, and award publication. The domain layer contains competition rules, scoring calculations, promotion logic, and audit concepts. The data access layer persists users, teams, rounds, submissions, scores, rankings, and export records.")
    add_figure(doc, images, "System Architecture Diagram", "Figure 11. System architecture diagram.")

    doc.add_heading("IV.3 Component and Package Diagram", level=2)
    doc.add_heading("IV.3.1 Component Diagram", level=2)
    add_figure(doc, images, "Page-1", "Figure 12. Component diagram.")
    doc.add_heading("IV.3.1 Package Diagram", level=2)
    add_table(doc, ["No", "Package", "Description"], [
        ["01", "identity", "Authentication, JWT issuance, password hashing, role and approval state management."],
        ["02", "event", "Event, round, category, criteria, mentor, judge assignment, and promotion rule management."],
        ["03", "team", "Team creation, member validation, and category registration."],
        ["04", "submission", "Round-based submission URLs and optional Git metadata synchronization."],
        ["05", "evaluation", "Criterion-level score capture, judge assignment enforcement, calibration support, and score audit."],
        ["06", "ranking", "Weighted aggregation, category/event ranking, promotion, elimination, and awards."],
        ["07", "reporting", "CSV/Excel exports, anonymized RBL datasets, and dashboard data."],
    ], [0.45, 1.55, 4.5])

    doc.add_heading("IV.4 Class diagram", level=2)
    add_para(doc, "The class model centers on Event, Round, Category, Team, TeamMember, Submission, Criteria, JudgeAssignment, Evaluation, ScoreDetail, RankingResult, Award, and AuditLog. The model separates score details from aggregated ranking results so SEAL can support both operational ranking and RBL research analysis.")
    add_figure(doc, images, "ClassDiagram", "Figure 13. Class diagram.")

    doc.add_heading("IV.5 Database Design", level=2)
    add_para(doc, "The relational database stores normalized competition data. Important tables include users, student_profiles, events, rounds, categories, teams, team_members, registrations, criteria_templates, event_criteria, judge_assignments, submissions, evaluations, score_details, rankings, eliminations, awards, notifications, and audit_logs.")
    add_figure(doc, images, "Database_Diagram_ERD", "Figure 14. Database ERD.")

    doc.add_heading("V. Implementation", level=1)
    doc.add_heading("V.1 Map architecture to the structure of the project", level=2)
    add_table(doc, ["Layer", "Suggested Project Module", "Responsibilities"], [
        ["Presentation", "web / api", "Screens and endpoints for account approval, events, teams, submissions, judging, ranking, exports, and dashboards."],
        ["Application", "services", "Use case orchestration, validation, transaction boundaries, and integration with GitHub/GitLab metadata APIs."],
        ["Domain", "domain", "Entities, value objects, scoring rules, promotion rules, elimination policy, and award logic."],
        ["Persistence", "repositories / data", "Database access for users, events, teams, submissions, evaluations, scores, rankings, and audit logs."],
        ["Infrastructure", "integrations / jobs", "JWT, email/notification delivery, CSV/Excel export, metadata synchronization, and scheduled ranking tasks."],
    ], [1.25, 1.7, 3.55])
    doc.add_heading("V.2 Map Class Diagram and Interaction Diagram to Code", level=2)
    add_para(doc, "Each sequence diagram maps to an application service method. For example, UC-02 maps to SubmissionService.submitProject, UC-03 maps to EvaluationService.scoreSubmission, and UC-04 maps to RankingService.generateRoundRanking. Controllers validate request permissions and delegate business decisions to services, while repositories encapsulate database access.")

    doc.add_heading("VI. Applying Alternative Architecture Patterns", level=1)
    doc.add_heading("VI.1 Applying Service-Oriented Architecture (SOA)", level=2)
    add_para(doc, "SEAL can evolve from a layered monolith into a service-oriented design by separating Identity Service, Event Service, Team Service, Submission Service, Evaluation Service, Ranking Service, Reporting Service, and Notification Service. This pattern is useful when different competition workflows need to scale independently or be maintained by separate teams.")
    add_bullets(doc, [
        "Identity Service owns user registration, account approval, roles, JWT, and guest judge accounts.",
        "Event Service owns event, round, category, criteria, mentor, and judge-assignment configuration.",
        "Submission and Evaluation Services own round artifacts and criterion-level scoring records.",
        "Ranking and Reporting Services own aggregation, promotion, awards, exports, and RBL datasets.",
    ])
    doc.add_heading("VI.2 Applying Service Discovery Pattern in the service-oriented architecture", level=2)
    add_para(doc, "If the system is deployed as multiple services, service discovery allows API gateway and backend services to locate available service instances dynamically. For the course project scope, a modular layered architecture is sufficient; service discovery is an optional future extension for larger deployments.")

    doc.add_heading("VII. Conclusion", level=1)
    add_para(doc, "The SEAL Hackathon Management System provides an end-to-end platform for managing Software Engineering hackathon events. By combining event configuration, team registration, submission management, judge evaluation, automated ranking, audit logging, and exportable research datasets, the system reduces manual coordination effort and improves transparency of competition results.")

    doc.save(OUT_DOCX)


def main() -> None:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    images = render_all_diagrams()
    build_report(images)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
